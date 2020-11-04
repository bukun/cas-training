

import requests
from bs4 import BeautifulSoup
import os
import docx
from docx import Document
from docx.shared import Inches


# 解析页面

# In[128]:


url = 'http://drr.ikcest.org/'
html = requests.get(url).content
soup = BeautifulSoup(html,'html.parser')
imgs_table = soup.find('table',{"class":"table"})
img=str(imgs_table.find('div',{"class":"col-sm-4"})).split('src="')[1].split('"')[0]
img_src='http://drr.ikcest.org'+img
img_title=imgs_table.find('div',{"class":"col-sm-8"}).text




img_name = 'xx_request.jpg'

with open(img_name,'wb')as f:
    response = requests.get(img_src).content
    f.write(response)
    f.close()


# 创建document对象，并向文档中添加文字，图片

# In[120]:


document = Document()
document.add_paragraph(img_title)
document.add_picture(img_name)


# 保存文档

# In[121]:


document.save('xx_docx.docx')


# 删除保存在本地的图片

# In[122]:


# os.remove(img_name)

