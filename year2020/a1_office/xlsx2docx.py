from openpyxl import load_workbook
from docx import Document
from xlsx_io import get_data

xlsx_obj = load_workbook('./demo_tmpl/e_mod.xlsx')
work_sheet = xlsx_obj.get_sheet_by_name('Sheet1')


def foo1():
    word = Document('./demo_tmpl/d_mod.docx')
    table1 = word.add_table(10, 7)
    # word.save('./demo_tmpl/d_mod.docx')

    for i in range(0, 10):
        for j in range(0, 7):
            table1.cell(i, j).text = str(work_sheet.cell(row=(i + 1), column=j + 1).value)
        # table1.cell(i, 1).text = str(ws.cell(row=(i + 1), column=2).value)

    word.save('./xx_d_mod.docx')


def foo2():
    tmpl = Document('./xx_d_mod.docx')
    paras = tmpl.paragraphs
    data_arr = get_data(None)
    for para in paras:
        print(para.text)
        para.text = para.text.replace('{a}', str(len(data_arr)))
        para.text = para.text.replace('{b}', str(sum(data_arr) / len(data_arr)))
        print(para.text)

    tmpl.save('xx_tmpl.docx')


if __name__ == '__main__':
    # foo1()
    foo2()
